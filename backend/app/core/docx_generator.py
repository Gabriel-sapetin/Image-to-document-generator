"""
Word document generation using python-docx with Grid Support
"""

import logging
from typing import List
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Page dimensions in inches
PAGE_DIMS = {
    'A4':     (8.27, 11.69),
    'LETTER': (8.5,  11.0),
    'SHORT':  (8.5,  11.0),
    'LONG':   (8.5,  13.0),
}

MARGIN_IN = 0.5  # inches


class WordGenerator:
    """Generate Word documents with configurable grid layout"""

    def generate(
        self,
        pages: List[List[dict]],
        output_path: str,
        title: str = "Image Collection",
        grid_cols: int = 2,
        page_size: str = 'A4'
    ) -> str:
        """
        Generate Word document with grid layout.
        
        Args:
            pages: Paginated image data
            output_path: Output file path
            title: Document title
            grid_cols: Number of columns per row
            page_size: 'A4', 'LETTER', 'SHORT', or 'LONG'
        """
        doc = Document()

        pw, ph = PAGE_DIMS.get(page_size.upper(), PAGE_DIMS['A4'])

        for section in doc.sections:
            section.page_width = Inches(pw)
            section.page_height = Inches(ph)
            section.top_margin = Inches(MARGIN_IN)
            section.bottom_margin = Inches(MARGIN_IN)
            section.left_margin = Inches(MARGIN_IN)
            section.right_margin = Inches(MARGIN_IN)

        # Title
        title_para = doc.add_paragraph(title)
        title_para.style = 'Heading 1'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        tp = doc.add_paragraph(f"Generated on {ts}")
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        usable_width = pw - MARGIN_IN * 2  # inches

        for page_num, page_images in enumerate(pages):
            if page_num > 0:
                doc.add_page_break()

            count = len(page_images)
            cols = min(grid_cols, count)
            rows = -(-count // cols)

            # Cell width per column minus small gap
            cell_w = (usable_width / cols) - 0.1

            for row_idx in range(rows):
                # Build a table row for each grid row
                start = row_idx * cols
                row_images = page_images[start:start + cols]

                table = doc.add_table(rows=1, cols=len(row_images))
                table.style = 'Table Grid'
                # Remove borders
                for cell in table.rows[0].cells:
                    for border in ['top', 'bottom', 'left', 'right']:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        border_el = OxmlElement(f'w:{border}')
                        border_el.set(qn('w:val'), 'none')
                        tcBorders.append(border_el)
                        tcPr.append(tcBorders)

                for col_idx, img_data in enumerate(row_images):
                    cell = table.rows[0].cells[col_idx]
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    try:
                        run.add_picture(img_data['filepath'], width=Inches(cell_w))
                    except Exception as e:
                        logger.error(f"Failed to insert {img_data.get('filepath')}: {e}")
                        para.add_run(f"[Image error: {e}]")

                # Small spacing between rows
                doc.add_paragraph()

        doc.save(output_path)
        logger.info(f"Word document saved: {output_path}")
        return output_path